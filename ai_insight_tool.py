import streamlit as st
import pandas as pd
import openai
import io
import docx
from docx import Document
from fpdf import FPDF
import xlsxwriter
import tempfile
import os
import re
import json

# ✅ Initialize OpenAI client
client = openai.OpenAI(
    api_key=st.secrets["openai_api_key"],
    project="proj_f0kWLTOK35ZQ5JKBmb2Y5Su9"
)

# ✅ UI: Title and Info
st.title("📊 Spreadsheet + Document Insight Bot")
st.markdown("""
Upload an **Excel (.xlsx)** or **Word (.docx)** file and ask natural language questions about its content.  
Your data is processed **in memory only** and never stored or transmitted.
""")

# ✅ Initialize session state
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "content_for_gpt" not in st.session_state:
    st.session_state.content_for_gpt = ""
if "last_answer" not in st.session_state:
    st.session_state.last_answer = ""
if "json_data" not in st.session_state:
    st.session_state.json_data = None
if "show_json" not in st.session_state:
    st.session_state.show_json = False

# ✅ File uploader
uploaded_file = st.file_uploader("Upload your file (.xlsx or .docx)", type=["xlsx", "docx"])

def is_excel(filename):
    return filename.lower().endswith(".xlsx")

def is_word(filename):
    return filename.lower().endswith(".docx")

MAX_ROWS = 500

if uploaded_file:
    try:
        file_name = uploaded_file.name

        if is_excel(file_name):
            df = pd.read_excel(uploaded_file)
            st.success("✅ Excel file uploaded successfully.")

            if len(df) > MAX_ROWS:
                st.warning(f"This file has {len(df)} rows. For best performance, only the first {MAX_ROWS} rows will be used.")
                df = df.head(MAX_ROWS)

            st.markdown("#### Preview (First Rows Analyzed)")
            st.dataframe(df)
            st.session_state.df = df
            st.session_state.content_for_gpt = df.to_string(index=False)

        elif is_word(file_name):
            doc = docx.Document(uploaded_file)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            st.success("✅ Word document uploaded successfully.")
            st.text_area("📄 Document Preview", value=full_text[:1000], height=200)
            st.session_state.content_for_gpt = full_text[:3000]

        else:
            st.warning("Unsupported file type.")
            st.session_state.content_for_gpt = ""

    except Exception as e:
        st.error(f"❌ Error reading file:\n\n{e}")

# ✅ Ask a question
if st.session_state.content_for_gpt:
    question = st.text_input("Ask a question about the content:")

    if question:
        with st.spinner("💡 Thinking..."):
            messages = [
                {"role": "system", "content": "You're an expert data and business assistant. When helpful, format structured data as a JSON array of objects and keep human-friendly summary text separate."},
                {"role": "user", "content": f"Here is the content to analyze:\n{st.session_state.content_for_gpt}"}
            ]
            for chat in st.session_state.chat_history:
                messages.append(chat)
            messages.append({"role": "user", "content": question})

            response = client.chat.completions.create(
                model="gpt-4o",
                messages=messages,
                temperature=0.2,
                max_tokens=1000
            )
            answer = response.choices[0].message.content
            st.session_state.chat_history.append({"role": "user", "content": question})
            st.session_state.chat_history.append({"role": "assistant", "content": answer})
            st.session_state.last_answer = answer

            human_friendly_text = re.sub(r"```json.*?```", "", answer, flags=re.DOTALL).strip()
            st.markdown("**💬 AI Response:**")
            st.markdown(human_friendly_text)

            try:
                match = re.search(r"```json\s*(.*?)```", answer, re.DOTALL)
                if match:
                    st.session_state.json_data = json.loads(match.group(1))
                else:
                    st.session_state.json_data = None
            except:
                st.session_state.json_data = None

# ✅ Download section if JSON extracted
if st.session_state.json_data:
    with st.expander("📦 View Extracted JSON Data"):
        st.download_button(
            "⬇️ Download JSON", 
            data=json.dumps(st.session_state.json_data, indent=2), 
            file_name="ai_response.json", 
            mime="application/json"
        )

    all_columns = list(st.session_state.json_data[0].keys()) if isinstance(st.session_state.json_data, list) else []
    selected_columns = st.multiselect("Select columns to include in download (from JSON):", options=all_columns, default=all_columns)

    if selected_columns:
        df_selected = pd.DataFrame(st.session_state.json_data)[selected_columns]

        # Excel Download
        excel_buffer = io.BytesIO()
        with pd.ExcelWriter(excel_buffer, engine="xlsxwriter") as writer:
            sheet = "AI Response"
            df_selected.to_excel(writer, index=False, sheet_name=sheet)
            workbook = writer.book
            worksheet = writer.sheets[sheet]
            header_format = workbook.add_format({'bold': True, 'bg_color': '#DCE6F1'})
            for col_num, value in enumerate(df_selected.columns.values):
                worksheet.write(0, col_num, value, header_format)
            worksheet.autofilter(0, 0, len(df_selected), len(df_selected.columns) - 1)
            for i, col in enumerate(df_selected.columns):
                width = max(df_selected[col].astype(str).map(len).max(), len(col)) + 2
                worksheet.set_column(i, i, width)
        excel_buffer.seek(0)
        st.download_button("⬇️ Download as .xlsx", data=excel_buffer, file_name="ai_response.xlsx")

        # Word Download
        word_buffer = io.BytesIO()
        doc = Document()
        table = doc.add_table(rows=1, cols=len(selected_columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(selected_columns):
            hdr_cells[i].text = col
        for index, row in df_selected.iterrows():
            row_cells = table.add_row().cells
            for i, col in enumerate(selected_columns):
                row_cells[i].text = str(row[col])
        doc.save(word_buffer)
        word_buffer.seek(0)
        st.download_button("⬇️ Download as .docx", data=word_buffer, file_name="ai_response.docx")

        # PDF Download
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=10)
        col_widths = [40] * len(selected_columns)
        row_height = 10
        for col in selected_columns:
            pdf.cell(col_widths[0], row_height, txt=col, border=1)
        pdf.ln()
        for _, row in df_selected.iterrows():
            for i, col in enumerate(selected_columns):
                text = str(row[col])[:60]
                pdf.cell(col_widths[i], row_height, txt=text, border=1)
            pdf.ln()
        pdf_buffer = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        pdf.output(pdf_buffer.name)
        with open(pdf_buffer.name, "rb") as f:
            st.download_button("⬇️ Download as .pdf", data=f, file_name="ai_response.pdf")
        os.unlink(pdf_buffer.name)

# ✅ AI Summary toggle
if st.session_state.content_for_gpt:
    if st.button("Summarize File with AI"):
        with st.spinner("Summarizing your data..."):
            summary_prompt = f"""
You're an AI business analyst. Give a short summary of this customer data. Focus only on the contents below — do not refer to previous questions. Your job is to:

- Highlight themes and trends in the customer notes
- Mention average opportunity score and spend if possible
- Suggest smart next steps for marketing or sales

Here’s the data:
{st.session_state.content_for_gpt}
            """
            summary_response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You're an expert data and business analyst. Be helpful, concise, and insightful."},
                    {"role": "user", "content": summary_prompt}
                ],
                temperature=0.3,
                max_tokens=700
            )
            summary = summary_response.choices[0].message.content
            st.markdown("### 🔍 AI Summary")
            st.info(summary)

# ✅ Optional: Display full chat log
if st.session_state.chat_history:
    with st.expander("🗂 View Chat History"):
        for i, msg in enumerate(st.session_state.chat_history):
            st.markdown(f"**{msg['role'].capitalize()}:** {msg['content']}")
